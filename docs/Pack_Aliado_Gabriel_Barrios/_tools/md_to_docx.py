#!/usr/bin/env python3
"""Convert Markdown files in pack md/ to docx/."""
import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PACK_ROOT = Path(__file__).resolve().parent.parent
MD_DIR = PACK_ROOT / "md"
DOCX_DIR = PACK_ROOT / "docx"


def set_cell_shading(cell, fill="1E2A5A"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            cell.text = row[j].strip() if j < len(row) else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if i == 0:
                set_cell_shading(cell)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph()


def md_to_docx(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                if re.match(r"^\|[\s\-:|]+\|$", lines[i]):
                    i += 1
                    continue
                cells = [c.strip() for c in lines[i].strip("|").split("|")]
                table_rows.append(cells)
                i += 1
            add_table(doc, table_rows)
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            if any(k in line for k in ("No sumes", "Regla", "importante")):
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "FFF3CD")
                pPr.append(shd)
            i += 1
            continue
        if line.startswith("---"):
            doc.add_paragraph("─" * 40)
            i += 1
            continue
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("md_file", nargs="?", help="Single md file or all if omitted")
    args = parser.parse_args()

    if args.md_file:
        md_path = Path(args.md_file)
        if not md_path.is_absolute():
            md_path = MD_DIR / md_path.name if not md_path.exists() else md_path
        out_path = DOCX_DIR / (md_path.stem + ".docx")
        md_to_docx(md_path, out_path)
        print(f"OK: {out_path}")
        return

    DOCX_DIR.mkdir(exist_ok=True)
    for md_path in sorted(MD_DIR.glob("*.md")):
        out_path = DOCX_DIR / (md_path.stem + ".docx")
        md_to_docx(md_path, out_path)
        print(f"OK: {out_path.name}")


if __name__ == "__main__":
    main()
